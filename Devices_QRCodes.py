import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
import os
import re
import subprocess
import platform

def modify_docx(file_path, start_number):
    try:
        # Load the document
        doc = Document(file_path)
        current_number = start_number
        replacement_count = 0

        # Iterate through each table in the document
        for table_index, table in enumerate(doc.tables):
            print(f"Processing table {table_index + 1}")
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    paras = cell.paragraphs
                    i = 0
                    while i < len(paras) - 2:
                        line1 = paras[i].text.strip()
                        line2 = paras[i+1].text.strip()
                        line3 = paras[i+2].text.strip()

                        print(f"Table {table_index + 1}, Row {row_index + 1}, Cell {cell_index + 1}:")
                        print(f"  Line 1: {line1}")
                        print(f"  Line 2: {line2}")
                        print(f"  Line 3: {line3}")

                        if line1 == "Property of" and line2 == "ABM" and re.match(r"^\d{4}$", line3):
                            full = f"Property of ABM {current_number}"
                            paras[i].text = full

                            # Set font size and type for "Property of"
                            for run in paras[i].runs:
                                run.font.size = Pt(11)
                                run.font.name = 'Calibri'

                            # Clear the text of the next two paragraphs to avoid affecting images
                            for run in paras[i+1].runs:
                                run.text = ""
                            for run in paras[i+2].runs:
                                run.text = ""

                            # Set font size and type for "ABM" and the number
                            for run in paras[i].runs:
                                if run.text.strip() == "ABM" or re.match(r"^\d{4}$", run.text.strip()):
                                    run.font.size = Pt(14)
                                    run.font.name = 'Calibri'

                            replacement_count += 1
                            if replacement_count % 2 == 0:
                                current_number += 1
                            i += 3
                        else:
                            i += 1

        # Save the modified document
        new_path = os.path.splitext(file_path)[0] + f"_from_{start_number}_modified.docx"
        doc.save(new_path)
        print(f"Document saved as: {new_path}")
        return new_path

    except Exception as e:
        print(f"An error occurred: {e}")
        raise

def open_file(filepath):
    if platform.system() == "Windows":
        os.startfile(filepath)
    elif platform.system() == "Darwin":  # macOS
        subprocess.call(["open", filepath])
    else:  # Linux
        subprocess.call(["xdg-open", filepath])

def run_app():
    def start_processing():
        file_path = file_entry.get()
        try:
            start_number = int(start_entry.get())
            new_path = modify_docx(file_path, start_number)
            messagebox.showinfo("Success", f"File saved as:\n{new_path}")
            open_file(new_path)
        except Exception as e:
            messagebox.showerror("Error", f"Something went wrong:\n{e}")

    def browse_file():
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)

    root = tk.Tk()
    root.title("ABM Modifier App")
    root.geometry("400x250")

    tk.Label(root, text="Select .docx file:").pack(pady=5)
    file_entry = tk.Entry(root, width=40)
    file_entry.pack()
    tk.Button(root, text="Browse", command=browse_file).pack(pady=5)

    tk.Label(root, text="Starting number:").pack()
    start_entry = tk.Entry(root)
    start_entry.pack()

    tk.Button(root, text="START", command=start_processing, bg="green", fg="white").pack(pady=10)

    root.mainloop()

if __name__ == "__main__":
    run_app()